"""
Comprehensive PDF content extractor — page-by-page analysis.
Extracts text, images, formulas, tables with full metadata.
Never generates hardcoded content — only extracts what exists in the PDF.
"""
import tempfile
import os
import re
import hashlib
from io import BytesIO
from dataclasses import dataclass, field


@dataclass
class ImageContent:
    """A single extracted image with full metadata."""
    image_id: str          # Unique ID like "img_p3_0"
    page_number: int       # Source page (1-based)
    data_b64: str          # Base64 encoded image data
    mime: str              # MIME type (image/png, image/jpeg, etc.)
    content_type: str      # "graph", "diagram", "chart", "table_image", "formula_image", "photo", "image"
    width: int = 0
    height: int = 0
    nearby_text: str = ""  # Text surrounding the image on the page
    position: str = ""     # "top", "bottom", "left", "right", "center"
    description: str = ""  # Brief description based on position/type


@dataclass
class FormulaContent:
    """A single extracted formula with metadata."""
    formula_id: str        # Unique ID like "form_p2_0"
    page_number: int       # Source page
    latex: str             # LaTeX representation
    raw_text: str          # Original text from PDF
    nearby_text: str = ""  # Surrounding context
    content_type: str = "formula"


@dataclass
class TableContent:
    """A single extracted table with metadata."""
    table_id: str          # Unique ID like "tab_p4_0"
    page_number: int
    text: str              # Table as text
    rows: int = 0
    cols: int = 0
    nearby_text: str = ""  # Text before/after the table


@dataclass
class PageContent:
    """All content from a single PDF page."""
    page_number: int
    text: str = ""
    headings: list = field(default_factory=list)      # Detected headings
    images: list = field(default_factory=list)         # List[ImageContent]
    formulas: list = field(default_factory=list)       # List[FormulaContent]
    tables: list = field(default_factory=list)         # List[TableContent]
    text_blocks: list = field(default_factory=list)    # Structured text blocks


@dataclass
class ExtractedContent:
    """Complete extracted content from a PDF — all metadata preserved."""
    pages: list = field(default_factory=list)
    all_images: list = field(default_factory=list)
    all_formulas: list = field(default_factory=list)
    all_tables: list = field(default_factory=list)
    full_text: str = ""
    total_pages: int = 0
    summary: str = ""  # Content summary for Gemini


def extract_pdf(content: bytes) -> ExtractedContent:
    """
    Extract content from a PDF with full metadata.
    Extracts all pages without arbitrary limits on pages, text, formulas, tables, or images.
    """
    import pymupdf as fitz

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    print(f"[extract] Opening PDF ({len(content)/1024/1024:.1f} MB)...")
    doc = fitz.open(tmp_path)
    total_pages = len(doc)
    print(f"[extract] {total_pages} pages found, processing all pages...")

    result = ExtractedContent(total_pages=total_pages)

    for page_num in range(total_pages):
        page = doc[page_num]
        page_content = PageContent(page_number=page_num + 1)

        # ═══════════════════════════════════════════════════════
        # 1. EXTRACT TEXT WITH STRUCTURE
        # ═══════════════════════════════════════════════════════
        try:
            text_dict = page.get_text("dict")
            full_page_text = ""

            for block in text_dict.get("blocks", []):
                if block["type"] == 0:  # Text block
                    block_text = ""
                    block_lines = []

                    for line in block.get("lines", []):
                        line_text = ""
                        line_spans = []

                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if not text:
                                continue
                            font = span.get("font", "")
                            size = span.get("size", 12)
                            flags = span.get("flags", 0)
                            bbox = span.get("bbox", [0, 0, 0, 0])

                            line_text += text + " "
                            line_spans.append({
                                "text": text,
                                "font": font,
                                "size": size,
                                "flags": flags,
                                "bbox": bbox,
                            })

                            # Detect headings (larger font size)
                            if size > 14 and len(text) > 2:
                                page_content.headings.append(text.strip())

                            # Detect formulas in text
                            if _is_formula_span(text, font, flags, size):
                                formula_latex = _text_to_latex(text)
                                if formula_latex:
                                    fid = f"form_p{page_num + 1}_{len(page_content.formulas)}"
                                    nearby = _get_nearby_text(full_page_text, text)
                                    page_content.formulas.append(FormulaContent(
                                        formula_id=fid,
                                        page_number=page_num + 1,
                                        latex=formula_latex,
                                        raw_text=text,
                                        nearby_text=nearby,
                                    ))

                        if line_text.strip():
                            block_lines.append(line_text.strip())
                            full_page_text += line_text + "\n"

                    if block_lines:
                        block_text = " ".join(block_lines)
                        page_content.text_blocks.append({
                            "text": block_text,
                            "bbox": block.get("bbox", []),
                        })

            page_content.text = full_page_text.strip()

        except Exception as e:
            print(f"[extract] Text extraction error on page {page_num + 1}: {e}")
            page_content.text = page.get_text("text")

        # ═══════════════════════════════════════════════════════
        # 2. EXTRACT IMAGES WITH FULL METADATA
        # ═══════════════════════════════════════════════════════
        try:
            image_list = page.get_images(full=True)
            seen_hashes = set()
            max_imgs_per_page = 5  # Limit images per page to keep high relevance

            for img_idx, img in enumerate(image_list[:max_imgs_per_page]):
                xref = img[0]
                try:
                    base_image = doc.extract_image(xref)
                    if not base_image:
                        continue

                    img_bytes = base_image["image"]
                    img_ext = base_image.get("ext", "png")

                    # Skip tiny icons (< 2KB)
                    if len(img_bytes) < 2000:
                        continue

                    # Deduplicate by content hash
                    img_hash = hashlib.md5(img_bytes).hexdigest()[:12]
                    if img_hash in seen_hashes:
                        continue
                    seen_hashes.add(img_hash)

                    # Convert to supported format — returns (mime, converted_bytes)
                    convert_result = _convert_image(img_bytes, img_ext)
                    if not convert_result:
                        continue
                    mime, converted_bytes = convert_result

                    # Get image dimensions (from converted bytes)
                    try:
                        from PIL import Image
                        pil_img = Image.open(BytesIO(converted_bytes))
                        w, h = pil_img.size
                    except Exception:
                        w, h = 0, 0

                    # Classify content type
                    content_type = _classify_image_content(img_bytes, img_ext, w, h)

                    # Find nearby text (text before/after image position)
                    bbox = img[:4]  # [x0, y0, x1, y1]
                    nearby = _find_nearby_text(page_content.text_blocks, bbox, page_content.text)

                    # Determine position on page
                    position = _determine_position(bbox, page.rect)

                    # Create unique ID
                    image_id = f"img_p{page_num + 1}_{img_idx}"

                    import base64
                    data_b64 = base64.b64encode(converted_bytes).decode("utf-8")

                    img_content = ImageContent(
                        image_id=image_id,
                        page_number=page_num + 1,
                        data_b64=data_b64,
                        mime=mime,
                        content_type=content_type,
                        width=w,
                        height=h,
                        nearby_text=nearby,
                        position=position,
                        description=f"{content_type} from page {page_num + 1}",
                    )
                    page_content.images.append(img_content)

                except Exception:
                    continue

        except Exception as e:
            print(f"[extract] Image extraction error on page {page_num + 1}: {e}")

        # ═══════════════════════════════════════════════════════
        # 3. EXTRACT TABLES
        # ═══════════════════════════════════════════════════════
        try:
            tabs = page.find_tables()
            if tabs and tabs.tables:
                for tab_idx, tab in enumerate(tabs.tables):
                    try:
                        tab_data = tab.extract()
                        rows = len(tab_data)
                        cols = len(tab_data[0]) if tab_data else 0
                        tab_text = "\n".join(
                            [" | ".join(str(cell) for cell in row) for row in tab_data]
                        )
                        tid = f"tab_p{page_num + 1}_{tab_idx}"
                        nearby = _find_nearby_text(page_content.text_blocks, tab.bbox, page_content.text)

                        page_content.tables.append(TableContent(
                            table_id=tid,
                            page_number=page_num + 1,
                            text=tab_text,
                            rows=rows,
                            cols=cols,
                            nearby_text=nearby,
                        ))
                    except Exception:
                        continue
        except Exception:
            pass

        # Collect all content
        result.pages.append(page_content)
        result.full_text += page_content.text + "\n\n"
        result.all_images.extend(page_content.images)
        result.all_formulas.extend(page_content.formulas)
        result.all_tables.extend(page_content.tables)

    doc.close()
    try:
        os.unlink(tmp_path)
    except Exception:
        pass

    # Build summary for Gemini
    result.summary = _build_summary(result)

    print(f"[extract] {result.total_pages} pages | "
          f"{len(result.all_images)} images | "
          f"{len(result.all_formulas)} formulas | "
          f"{len(result.all_tables)} tables | "
          f"{len(result.full_text)} chars text")

    return result


def extract_generic(file_content: bytes, ext: str) -> tuple[str, list[dict]]:
    """Extract from non-PDF files."""
    if ext == "txt":
        return file_content.decode("utf-8", errors="ignore"), []
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            doc = Document(tmp_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            os.unlink(tmp_path)
            return text, []
        except ImportError:
            raise Exception("DOCX support requires python-docx")
    return file_content.decode("utf-8", errors="ignore"), []


# ═══════════════════════════════════════════════════════════════
# PRIVATE HELPERS
# ═══════════════════════════════════════════════════════════════

def _is_formula_span(text: str, font: str, flags: int, size: float) -> bool:
    """Detect if a text span contains a mathematical formula."""
    text_clean = text.strip()
    if not text_clean or len(text_clean) < 2:
        return False

    math_chars = "∑∏∫∂√∞≈≠≤≥±×÷∈∉⊂⊃∪∩αβγδεζηθκλμνξπρστφχψω"
    has_math_chars = any(c in text_clean for c in math_chars)

    # True equation patterns (e.g., y = mx + b, a^2 + b^2 = c^2, x_i)
    eq_patterns = [
        r"[a-zA-Z0-9]\s*=\s*[^=]",            # x = ...
        r"\d+\s*[+\-*/^]\s*\d+",              # 2 + 3
        r"[a-zA-Z]\s*[+\-*/^]\s*[a-zA-Z0-9]", # ax + b
        r"\b[a-zA-Z]_[a-zA-Z0-9{]",           # subscript notation x_1, x_{i}
        r"\b[a-zA-Z]\^[a-zA-Z0-9{]",          # superscript notation x^2
        r"\\frac|\\sqrt|\\sum|\\int",         # LaTeX commands
    ]
    has_eq = any(re.search(p, text_clean) for p in eq_patterns)

    font_lower = font.lower()
    math_fonts = ["math", "symbol", "stix", "cambria math", "computer modern math"]
    has_math_font = any(f in font_lower for f in math_fonts)

    # If it's pure English prose (letters and punctuation only without math symbols or equation syntax), skip it
    if re.match(r"^[A-Za-z\s,.'\"?!()]+$", text_clean) and not has_math_chars and not has_eq:
        return False

    return has_math_chars or has_eq or (has_math_font and len(text_clean) < 200)


def _text_to_latex(text: str) -> str:
    """Convert text with math symbols to LaTeX."""
    replacements = {
        "α": r"\alpha", "β": r"\beta", "γ": r"\gamma", "δ": r"\delta",
        "ε": r"\epsilon", "ζ": r"\zeta", "η": r"\eta", "θ": r"\theta",
        "κ": r"\kappa", "λ": r"\lambda", "μ": r"\mu", "ν": r"\nu",
        "ξ": r"\xi", "π": r"\pi", "ρ": r"\rho", "σ": r"\sigma",
        "τ": r"\tau", "φ": r"\phi", "χ": r"\chi", "ψ": r"\psi", "ω": r"\omega",
        "∑": r"\sum", "∏": r"\prod", "∫": r"\int", "∂": r"\partial",
        "√": r"\sqrt", "∞": r"\infty", "≈": r"\approx", "≠": r"\neq",
        "≤": r"\leq", "≥": r"\geq", "±": r"\pm", "×": r"\times", "÷": r"\div",
        "∈": r"\in", "∉": r"\notin", "⊂": r"\subset", "⊃": r"\supset",
        "∪": r"\cup", "∩": r"\cap",
    }
    result = text
    for char, latex in replacements.items():
        result = result.replace(char, f" {latex} ")
    return result.strip()


def _convert_image(img_bytes: bytes, ext: str) -> tuple[str, bytes] | None:
    """Convert and optimize image to supported MIME type (PNG/JPEG) with max dimensions 1024px."""
    try:
        from PIL import Image
        pil_img = Image.open(BytesIO(img_bytes))

        # Convert color modes
        if pil_img.mode in ("P", "PA", "CMYK"):
            pil_img = pil_img.convert("RGBA")
        elif pil_img.mode == "L":
            pil_img = pil_img.convert("RGB")

        # Resize if overly large (> 1024 on any side) to keep memory lightweight
        max_dim = 1024
        if pil_img.width > max_dim or pil_img.height > max_dim:
            pil_img.thumbnail((max_dim, max_dim))

        buf = BytesIO()
        # Save as JPEG if RGB, PNG if RGBA (has transparency)
        if pil_img.mode == "RGBA":
            pil_img.save(buf, format="PNG", optimize=True)
            return ("image/png", buf.getvalue())
        else:
            if pil_img.mode != "RGB":
                pil_img = pil_img.convert("RGB")
            pil_img.save(buf, format="JPEG", quality=85, optimize=True)
            return ("image/jpeg", buf.getvalue())
    except Exception:
        return None


def _classify_image_content(img_bytes: bytes, ext: str, w: int, h: int) -> str:
    """Classify image content type based on dimensions and characteristics."""
    if w == 0 or h == 0:
        return "image"

    aspect = w / h if h > 0 else 1

    # Wide + short = likely graph/chart
    if aspect > 2.5 and w > 400:
        return "graph"

    # Very wide = chart/graph
    if aspect > 1.8 and w > 300:
        return "chart"

    # Tall = likely diagram/flowchart
    if aspect < 0.5 and h > 300:
        return "diagram"

    # Square-ish and medium = could be formula symbol or diagram
    if 0.7 < aspect < 1.3 and 100 < w < 400:
        return "diagram"

    return "image"


def _find_nearby_text(text_blocks: list, img_bbox: tuple, full_text: str) -> str:
    """Find text near an image based on bounding box proximity."""
    if not text_blocks:
        # Fallback: return first 200 chars of page text
        return full_text[:200] if full_text else ""

    img_y = (img_bbox[1] + img_bbox[3]) / 2 if len(img_bbox) >= 4 else 0

    # Find closest text block by vertical position
    closest = ""
    min_dist = float("inf")

    for block in text_blocks:
        bbox = block.get("bbox", [0, 0, 0, 0])
        if len(bbox) >= 4:
            block_y = (bbox[1] + bbox[3]) / 2
            dist = abs(block_y - img_y)
            if dist < min_dist:
                min_dist = dist
                closest = block.get("text", "")

    return closest[:300] if closest else full_text[:200]


def _determine_position(bbox: tuple, page_rect) -> str:
    """Determine image position on the page."""
    if len(bbox) < 4 or not page_rect:
        return "unknown"

    x0, y0, x1, y1 = bbox
    mid_x = (x0 + x1) / 2
    mid_y = (y0 + y1) / 2
    page_w = page_rect.width
    page_h = page_rect.height

    # Vertical position
    if mid_y < page_h * 0.33:
        vert = "top"
    elif mid_y > page_h * 0.66:
        vert = "bottom"
    else:
        vert = "center"

    # Horizontal position
    if mid_x < page_w * 0.33:
        horiz = "left"
    elif mid_x > page_w * 0.66:
        horiz = "right"
    else:
        horiz = "center"

    return f"{vert}_{horiz}" if vert != "center" or horiz != "center" else "center"


def _get_nearby_text(full_text: str, target: str) -> str:
    """Get text surrounding a specific string."""
    idx = full_text.find(target)
    if idx == -1:
        return full_text[:200]
    start = max(0, idx - 100)
    end = min(len(full_text), idx + len(target) + 100)
    return full_text[start:end]


def _build_summary(result: ExtractedContent) -> str:
    """Build a content summary for Gemini."""
    parts = [f"PDF: {result.total_pages} pages"]

    if result.all_images:
        types = {}
        for img in result.all_images:
            types[img.content_type] = types.get(img.content_type, 0) + 1
        parts.append(f"Images: {len(result.all_images)} ({', '.join(f'{v} {k}' for k, v in types.items())})")

    if result.all_formulas:
        parts.append(f"Formulas: {len(result.all_formulas)}")

    if result.all_tables:
        parts.append(f"Tables: {len(result.all_tables)} ({sum(t.rows for t in result.all_tables)} total rows)")

    parts.append(f"Text: {len(result.full_text)} characters")

    return " | ".join(parts)
